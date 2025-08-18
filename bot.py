# -*- coding: utf-8 -*-
"""
CV Telegram Bot — HTML/CSS + DocRaptor + DOCX (Render-Ready)

• Stack: python-telegram-bot v21 (async), SQLite, Jinja2 (HTML), DocRaptor (PDF/PNG), docxtpl (DOCX fallback).
• Deploy: Render (polling + /health mini server).
• Assets:
    - HTML templates: assets/html/<Template>_<lang>.html  (مثال: Navy_ar.html, Navy_en.html)
    - CSS مشترك اختياري: assets/html/base.css  (سيتم inlining تلقائيًا)
• PDF/PNG عبر DocRaptor: ضع DOCRAPTOR_API_KEY في المتغيرات البيئية.

ENV (Render):
BOT_TOKEN=xxxxxxxxxxxxxxxxxxxxxxxxxxxx
DB_PATH=/var/data/bot.db
PORT=10000           # Render يمرره تلقائيًا، هذا احتياط
OWNER_USERNAME=youruser   # VIP دائم للمالك (اختياري)
OWNER_ID=0                # بديل لتحديد VIP دائم (اختياري)
PAYLINK_UPGRADE_URL=https://your-pay-page (اختياري)
DOCRAPTOR_API_KEY=dp_xxxxxxxxxxxxxxxxx  (مطلوب للمعاينة/‏PDF)
ENABLE_PDF=0  # لتحويل DOCX->PDF عبر LibreOffice (اختياري جدًا)
"""

import asyncio
import json
import logging
import os
import re
import shutil
import sqlite3
import textwrap
from datetime import datetime, date
from pathlib import Path

from aiohttp import web
from jinja2 import Template
import httpx

from telegram import (
    Update,
    InlineKeyboardMarkup, InlineKeyboardButton,
    InputFile, BotCommand,
)
from telegram.constants import ChatAction
from telegram.ext import (
    Application, CommandHandler, CallbackQueryHandler,
    MessageHandler, ConversationHandler, ContextTypes, filters
)

# ---------------- Logging ----------------
logging.basicConfig(
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
    level=logging.INFO,
)
log = logging.getLogger("cvbot")

# ---------------- Config ----------------
DB_PATH = os.getenv("DB_PATH", "/var/data/bot.db")
OWNER_USERNAME = os.getenv("OWNER_USERNAME", "")
OWNER_ID = int(os.getenv("OWNER_ID", "0") or "0")
PAYLINK_UPGRADE_URL = os.getenv("PAYLINK_UPGRADE_URL", "")
ENABLE_PDF = os.getenv("ENABLE_PDF", "0") == "1"
DOCRAPTOR_API_KEY = os.getenv("DOCRAPTOR_API_KEY", "")
PORT = int(os.getenv("PORT", os.getenv("RENDER_PORT", "10000")))

TEMPLATES_DIR = Path("assets/templates")     # DOCX templates (اختياري)
HTML_TEMPLATES_DIR = Path("assets/html")     # HTML/CSS templates (موصى بها)
EXPORTS_DIR = Path(os.getenv("EXPORTS_DIR", "/var/data/exports"))
try:
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
except Exception:
    EXPORTS_DIR = Path("./exports")
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

# ---------------- Conversation States ----------------
(
    ASK_LANG, ASK_TPL, ASK_NAME, ASK_TITLE, ASK_PHONE, ASK_EMAIL, ASK_CITY, ASK_LINKS, ASK_SUMMARY,
    MENU, EXP_ROLE, EXP_COMPANY, EXP_START, EXP_END, EXP_BULLETS,
    EDU_DEGREE, EDU_MAJOR, EDU_SCHOOL, EDU_YEAR,
    SKILLS_SET, CONFIRM_EXPORT,
) = range(21)

TEMPLATES_INDEX = {
    # خلّي HTML القوالب اللي أضفتها في assets/html بنفس الأسماء Navy/Modern
    "ar": [("Navy", "احترافي (شريط جانبي أزرق)"), ("Modern", "حديث"), ("ATS", "مطابق ATS"), ("Minimal", "بسيط"), ("Elegant", "أنيق")],
    "en": [("Navy", "Professional (Navy Sidebar)"), ("Modern", "Modern"), ("ATS", "ATS"), ("Minimal", "Minimal"), ("Elegant", "Elegant")],
}

# ---------------- DB ----------------
class DB:
    def __init__(self, path: str):
        self.path = path
        try:
            Path(self.path).parent.mkdir(parents=True, exist_ok=True)
        except Exception:
            self.path = str(Path("./var_data/bot.db"))
            Path(self.path).parent.mkdir(parents=True, exist_ok=True)
        self._init()

    def _conn(self): return sqlite3.connect(self.path)

    def _init(self):
        con = self._conn(); cur = con.cursor()
        cur.executescript("""
        CREATE TABLE IF NOT EXISTS users(
          user_id INTEGER PRIMARY KEY, lang TEXT DEFAULT 'ar', vip INTEGER DEFAULT 0, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS cv_profile(
          id INTEGER PRIMARY KEY AUTOINCREMENT,
          user_id INTEGER, title TEXT, full_name TEXT,
          phone TEXT, email TEXT, city TEXT, links TEXT,
          summary TEXT, template TEXT DEFAULT 'Navy', lang TEXT DEFAULT 'ar',
          created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP, updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );
        CREATE TABLE IF NOT EXISTS cv_experience(
          id INTEGER PRIMARY KEY AUTOINCREMENT,
          profile_id INTEGER, company TEXT, role TEXT, start_date TEXT, end_date TEXT, bullets TEXT
        );
        CREATE TABLE IF NOT EXISTS cv_education(
          id INTEGER PRIMARY KEY AUTOINCREMENT,
          profile_id INTEGER, degree TEXT, major TEXT, school TEXT, year TEXT
        );
        CREATE TABLE IF NOT EXISTS cv_skills(
          id INTEGER PRIMARY KEY AUTOINCREMENT, profile_id INTEGER, skills TEXT
        );
        CREATE TABLE IF NOT EXISTS cv_once(user_id INTEGER PRIMARY KEY, used INTEGER DEFAULT 0);
        CREATE TABLE IF NOT EXISTS cv_quota(user_id INTEGER PRIMARY KEY, daily_used INTEGER DEFAULT 0, last_reset DATE);
        """)
        con.commit(); con.close()

    # users / vip
    def ensure_user(self, user_id: int, lang: str = "ar"):
        con=self._conn(); cur=con.cursor()
        cur.execute("SELECT 1 FROM users WHERE user_id=?", (user_id,))
        if not cur.fetchone():
            cur.execute("INSERT INTO users(user_id,lang,vip) VALUES(?,?,0)", (user_id, lang))
            con.commit()
        con.close()

    def is_vip(self, user_id: int) -> bool:
        con=self._conn(); cur=con.cursor()
        cur.execute("SELECT vip FROM users WHERE user_id=?", (user_id,))
        row=cur.fetchone(); con.close()
        return bool(row and row[0])

    def set_vip(self, user_id: int, vip: int):
        con=self._conn(); cur=con.cursor()
        cur.execute("UPDATE users SET vip=? WHERE user_id=?", (vip, user_id))
        con.commit(); con.close()

    # free-once (مدى الحياة)
    def free_once_available(self, user_id: int) -> bool:
        con=self._conn(); cur=con.cursor()
        cur.execute("SELECT used FROM cv_once WHERE user_id=?", (user_id,))
        row=cur.fetchone(); con.close()
        return (row is None) or (row[0]==0)

    def mark_free_once_used(self, user_id: int):
        con=self._conn(); cur=con.cursor()
        cur.execute("INSERT INTO cv_once(user_id,used) VALUES(?,1) ON CONFLICT(user_id) DO UPDATE SET used=1", (user_id,))
        con.commit(); con.close()

    # profile blocks
    def new_profile(self, user_id: int, lang: str, template: str) -> int:
        con=self._conn(); cur=con.cursor()
        cur.execute("INSERT INTO cv_profile(user_id,lang,template) VALUES(?,?,?)", (user_id,lang,template))
        pid=cur.lastrowid; con.commit(); con.close(); return pid

    def update_profile(self, pid: int, **fields):
        if not fields: return
        con=self._conn(); cur=con.cursor()
        keys=", ".join([f"{k}=?" for k in fields.keys()])
        cur.execute(f"UPDATE cv_profile SET {keys}, updated_at=CURRENT_TIMESTAMP WHERE id=?", (*fields.values(), pid))
        con.commit(); con.close()

    def add_experience(self, pid:int, company:str, role:str, start_date:str, end_date:str, bullets:list[str]):
        con=self._conn(); cur=con.cursor()
        cur.execute("INSERT INTO cv_experience(profile_id,company,role,start_date,end_date,bullets) VALUES(?,?,?,?,?,?)",
                    (pid,company,role,start_date,end_date,json.dumps(bullets, ensure_ascii=False)))
        con.commit(); con.close()

    def add_education(self, pid:int, degree:str, major:str, school:str, year:str):
        con=self._conn(); cur=con.cursor()
        cur.execute("INSERT INTO cv_education(profile_id,degree,major,school,year) VALUES(?,?,?,?,?)",
                    (pid,degree,major,school,year))
        con.commit(); con.close()

    def set_skills(self, pid:int, skills_str:str):
        con=self._conn(); cur=con.cursor()
        cur.execute("SELECT 1 FROM cv_skills WHERE profile_id=?", (pid,))
        if cur.fetchone():
            cur.execute("UPDATE cv_skills SET skills=? WHERE profile_id=?", (skills_str,pid))
        else:
            cur.execute("INSERT INTO cv_skills(profile_id,skills) VALUES(?,?)",(pid,skills_str))
        con.commit(); con.close()

    def fetch_full_profile(self, pid:int):
        con=self._conn(); cur=con.cursor()
        cur.execute("SELECT * FROM cv_profile WHERE id=?", (pid,))
        row=cur.fetchone()
        if not row: con.close(); return None
        cols=[d[0] for d in cur.description]
        profile=dict(zip(cols,row))

        cur.execute("SELECT company,role,start_date,end_date,bullets FROM cv_experience WHERE profile_id=?", (pid,))
        exps=[{"company":r[0],"role":r[1],"start_date":r[2],"end_date":r[3],
               "bullets": json.loads(r[4]) if r[4] else []} for r in cur.fetchall()]

        cur.execute("SELECT degree,major,school,year FROM cv_education WHERE profile_id=?", (pid,))
        edus=[{"degree":r[0],"major":r[1],"school":r[2],"year":r[3]} for r in cur.fetchall()]

        cur.execute("SELECT skills FROM cv_skills WHERE profile_id=?", (pid,))
        s=cur.fetchone(); skills=s[0] if s else ""
        con.close()
        return profile, exps, edus, skills

db = DB(DB_PATH)

# -------------- Helpers/Owner --------------
def user_is_owner(u) -> bool:
    try:
        if OWNER_ID and getattr(u,"id",None)==OWNER_ID: return True
    except Exception:
        pass
    if OWNER_USERNAME:
        uname=(getattr(u,"username","") or "").lower()
        if uname == OWNER_USERNAME.lower(): return True
    return False

def _safe(s:str|None)->str: return s or ""

# -------------- DOCX fallback (optional) --------------
from docxtpl import DocxTemplate
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document=None

def render_docx_for_profile(pid:int, db:DB)->Path:
    data=db.fetch_full_profile(pid)
    if not data: raise RuntimeError("Profile not found")
    profile, exps, edus, skills = data
    lang=profile.get("lang","ar"); tpl_slug=profile.get("template","Navy")

    ctx={
        "full_name":_safe(profile.get("full_name")),
        "title":_safe(profile.get("title")),
        "phone":_safe(profile.get("phone")),
        "email":_safe(profile.get("email")),
        "city":_safe(profile.get("city")),
        "links":_safe(profile.get("links")),
        "summary":_safe(profile.get("summary")),
        "experiences":exps, "education":edus, "skills":skills,
        "skills_list":[s.strip() for s in (skills or "").replace("؛",",").split(",") if s.strip()],
    }
    out_path=EXPORTS_DIR/f"cv_{pid}_{lang}.docx"
    tpl_path=TEMPLATES_DIR/f"{tpl_slug}_{lang}.docx"
    if tpl_path.exists():
        doc=DocxTemplate(tpl_path); doc.render(ctx); doc.save(out_path); return out_path

    # Auto simple DOCX (افتراضي)
    if Document is None: raise RuntimeError("No DOCX engine available")

    def shade_cell(cell,color_hex:str):
        tcPr=cell._tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),color_hex)
        tcPr.append(shd)

    docx=Document()
    for s in docx.sections:
        s.top_margin=s.bottom_margin=Inches(0.4); s.left_margin=s.right_margin=Inches(0.4)

    table=docx.add_table(rows=1,cols=2); table.autofit=False
    left,right=table.rows[0].cells
    table.columns[0].width=Inches(2.2); table.columns[1].width=Inches(4.8)
    NAVY=RGBColor(31,58,95); WHITE=RGBColor(255,255,255)
    shade_cell(left,'1f3a5f')

    def add_left_h(t):
        p=left.add_paragraph(); r=p.add_run(t); r.font.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE; p.space_after=Pt(2)
    def add_left_line(t):
        p=left.add_paragraph(); r=p.add_run(t); r.font.size=Pt(9); r.font.color.rgb=WHITE; p.space_after=Pt(1)

    add_left_h('التواصل' if lang=='ar' else 'Contact')
    for item in [ctx['phone'],ctx['email'],ctx['city'],ctx['links']]:
        if item: add_left_line(item)
    left.add_paragraph().space_after=Pt(6)

    add_left_h('التعليم' if lang=='ar' else 'Education')
    for ed in edus:
        add_left_line(f"{ed.get('degree','')} — {ed.get('school','')}")
        if ed.get('year'): add_left_line(str(ed.get('year')))
    left.add_paragraph().space_after=Pt(6)

    add_left_h('المهارات' if lang=='ar' else 'Skills')
    for s in ctx["skills_list"] or [ctx["skills"]]:
        if s: add_left_line(f"• {s}")

    p=right.add_paragraph()
    r=p.add_run(ctx['full_name']); r.font.size=Pt(20); r.font.bold=True; r.font.color.rgb=NAVY
    if ctx['title']:
        p.add_run("\n")   # ← مهم: سطر واحد، لا تكسره
        t=p.add_run(ctx['title']); t.font.size=Pt(12)
    right.add_paragraph()

    h=right.add_paragraph('الملخص' if lang=='ar' else 'Summary'); h.runs[0].font.bold=True; h.runs[0].font.size=Pt(12)
    if ctx['summary']:
        for line in textwrap.wrap(ctx['summary'],width=120):
            rp=right.add_paragraph(line); rp.paragraph_format.space_after=Pt(2)
    right.add_paragraph()

    h=right.add_paragraph('الخبرات' if lang=='ar' else 'Work Experience'); h.runs[0].font.bold=True; h.runs[0].font.size=Pt(12)
    for e in exps:
        p=right.add_paragraph(); rr=p.add_run(f"{e.get('role','')} — {e.get('company','')}"); rr.font.bold=True
        if e.get('start_date') or e.get('end_date'): p.add_run(f" ({e.get('start_date','')} - {e.get('end_date','')})")
        for b in e.get('bullets',[])[:6]:
            bp=right.add_paragraph(f"• {b}"); bp.paragraph_format.space_after=Pt(0)
    docx.add_paragraph()

    docx.save(out_path); return out_path

def try_convert_to_pdf(docx_path:Path)->Path|None:
    if not ENABLE_PDF: return None
    lo=shutil.which("libreoffice") or shutil.which("soffice")
    if not lo:
        log.warning("LibreOffice not found; skipping PDF convert")
        return None
    import subprocess
    try:
        subprocess.run([lo,"--headless","--convert-to","pdf","--outdir",str(docx_path.parent),str(docx_path)],
                       check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        out=docx_path.with_suffix(".pdf")
        return out if out.exists() else None
    except Exception as e:
        log.exception("PDF convert failed: %s", e)
        return None

# -------------- HTML/CSS rendering + DocRaptor --------------
def _inline_local_css(html:str, base_dir:Path)->str:
    """
    يستبدل <link rel="stylesheet" href="*.css"> بمحتوى CSS داخل <style>.
    """
    def repl(match):
        href=match.group(1)
        css_path=(base_dir / href).resolve()
        try:
            css=css_path.read_text(encoding="utf-8")
            return f"<style>\n{css}\n</style>"
        except Exception:
            return match.group(0)
    return re.sub(r'<link\s+[^>]*href=["\']([^"\']+\.css)["\'][^>]*>', repl, html, flags=re.I)

def render_html_for_profile(pid:int, db:DB)->str:
    data=db.fetch_full_profile(pid)
    if not data: raise RuntimeError("Profile not found")
    profile, exps, edus, skills=data
    lang=profile.get("lang","ar"); tpl_slug=profile.get("template","Navy")

    skills_list=[s.strip() for s in (skills or "").replace("؛",",").split(",") if s.strip()]
    ctx={
        "full_name":_safe(profile.get("full_name")),
        "title":_safe(profile.get("title")),
        "phone":_safe(profile.get("phone")),
        "email":_safe(profile.get("email")),
        "city":_safe(profile.get("city")),
        "links":_safe(profile.get("links")),
        "summary":_safe(profile.get("summary")),
        "experiences":exps, "education":edus,
        "skills":skills, "skills_list":skills_list,
        "photo_data_uri":"",  # إضافة لاحقة عند دعم الصور
    }

    tpl_path = HTML_TEMPLATES_DIR / f"{tpl_slug}_{lang}.html"
    if not tpl_path.exists():
        # Fallback HTML بسيط إذا القالب ناقص
        return f"""<!doctype html><meta charset="utf-8">
        <title>{ctx['full_name']}</title>
        <h1 style="font-family:Arial">{ctx['full_name']} — {ctx['title']}</h1>"""

    html_src = tpl_path.read_text(encoding="utf-8")
    html = Template(html_src).render(**ctx)
    # inline CSS (مثل base.css)
    html = _inline_local_css(html, tpl_path.parent)
    return html

async def docraptor_convert(html:str, kind:str="pdf")->bytes:
    """
    kind: 'pdf' أو 'png' (معاينة). يحتاج DOCRAPTOR_API_KEY.
    """
    if not DOCRAPTOR_API_KEY:
        raise RuntimeError("DOCRAPTOR_API_KEY is missing")
    payload = {"doc": {
        "test": False,            # True لو تبغى watermark مجاني
        "document_type": kind,    # 'pdf' أو 'png'
        "name": f"cv.{kind}",
        "document_content": html
    }}
    async with httpx.AsyncClient(timeout=120) as client:
        r = await client.post("https://api.docraptor.com/docs", auth=(DOCRAPTOR_API_KEY,""), json=payload)
        r.raise_for_status()
        return r.content

# -------------- Bot Handlers --------------
async def set_my_commands(app: Application):
    await app.bot.set_my_commands([
        BotCommand("start","ابدأ / Start"),
        BotCommand("cv","إنشاء/تعديل السيرة"),
        BotCommand("upgrade","الترقية إلى VIP"),
        BotCommand("help","مساعدة"),
    ])

async def start(update:Update, context:ContextTypes.DEFAULT_TYPE):
    u=update.effective_user; db.ensure_user(u.id)
    if user_is_owner(u): db.set_vip(u.id,1)
    await update.effective_message.reply_text(
        "أهلًا! هذا بوت إنشاء سيرة (HTML/CSS + PDF احترافي).\n"
        "أرسل /cv للبدء."
    )

async def help_cmd(update:Update, context:ContextTypes.DEFAULT_TYPE):
    await update.effective_message.reply_text("/cv للبدء • /upgrade للترقية • تواصل: @" + (OWNER_USERNAME or "admin"))

async def upgrade_cmd(update:Update, context:ContextTypes.DEFAULT_TYPE):
    if PAYLINK_UPGRADE_URL:
        await update.effective_message.reply_text(f"للترقية إلى VIP: {PAYLINK_UPGRADE_URL}")
    else:
        await update.effective_message.reply_text("ضع PAYLINK_UPGRADE_URL في المتغيرات البيئية.")

# --- CV flow ---
async def cv_entry(update:Update, context:ContextTypes.DEFAULT_TYPE):
    await update.effective_chat.send_action(ChatAction.TYPING)
    context.user_data["cv"]={}
    kb=[[InlineKeyboardButton("عربي",callback_data="cv:lang:ar"),
         InlineKeyboardButton("English",callback_data="cv:lang:en")]]
    await update.effective_message.reply_text("اختر لغة السيرة:", reply_markup=InlineKeyboardMarkup(kb))
    return ASK_LANG

async def cv_set_lang(update:Update, context:ContextTypes.DEFAULT_TYPE):
    q=update.callback_query; await q.answer()
    lang=q.data.split(":")[-1]; context.user_data["cv"]["lang"]=lang
    tpl_buttons=[[InlineKeyboardButton(name, callback_data=f"cv:tpl:{slug}")]
                 for slug,name in TEMPLATES_INDEX[lang]]
    await q.edit_message_text("اختر القالب:", reply_markup=InlineKeyboardMarkup(tpl_buttons))
    return ASK_TPL

async def cv_set_tpl(update:Update, context:ContextTypes.DEFAULT_TYPE):
    q=update.callback_query; await q.answer()
    tpl_slug=q.data.split(":")[-1]; context.user_data["cv"]["template"]=tpl_slug
    await q.edit_message_text(f"تم اختيار قالب: {tpl_slug}\nأرسل اسمك الكامل:")
    return ASK_NAME

async def cv_name(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["cv"]["full_name"]=update.message.text.strip()
    await update.message.reply_text("المسمى الوظيفي المستهدف:")
    return ASK_TITLE

async def cv_title(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["cv"]["title"]=update.message.text.strip()
    await update.message.reply_text("رقم الجوال:")
    return ASK_PHONE

async def cv_phone(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["cv"]["phone"]=update.message.text.strip()
    await update.message.reply_text("البريد الإلكتروني:")
    return ASK_EMAIL

async def cv_email(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["cv"]["email"]=update.message.text.strip()
    await update.message.reply_text("المدينة:")
    return ASK_CITY

async def cv_city(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["cv"]["city"]=update.message.text.strip()
    await update.message.reply_text("روابطك (LinkedIn/GitHub) مفصولة بفواصل أو اكتب - لا يوجد -:")
    return ASK_LINKS

async def cv_links(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["cv"]["links"]=update.message.text.strip()
    await update.message.reply_text("اكتب ملخصًا قصيرًا (3-4 أسطر):")
    return ASK_SUMMARY

async def cv_summary(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["cv"]["summary"]=update.message.text.strip()
    u=update.effective_user; cv=context.user_data["cv"]
    db.ensure_user(u.id, cv.get("lang","ar"))
    pid=db.new_profile(u.id, cv.get("lang","ar"), cv.get("template","Navy"))
    db.update_profile(pid,
        full_name=cv.get("full_name"), title=cv.get("title"),
        phone=cv.get("phone"), email=cv.get("email"), city=cv.get("city"),
        links=cv.get("links"), summary=cv.get("summary"),
    )
    context.user_data["cv"]["pid"]=pid
    await show_menu(update, context, pid)
    return MENU

async def show_menu(update_or_q, context:ContextTypes.DEFAULT_TYPE, pid:int):
    txt=("القائمة الرئيسية:\n"
         "• إضافة خبرة\n• إضافة تعليم\n• تعيين المهارات\n• معاينة/تصدير")
    kb=[
        [InlineKeyboardButton("➕ إضافة خبرة", callback_data=f"cv:menu:addexp:{pid}")],
        [InlineKeyboardButton("🎓 إضافة تعليم", callback_data=f"cv:menu:addedu:{pid}")],
        [InlineKeyboardButton("🧩 تعيين المهارات", callback_data=f"cv:menu:skills:{pid}")],
        [InlineKeyboardButton("📤 معاينة/تصدير", callback_data=f"cv:menu:export:{pid}")],
    ]
    if isinstance(update_or_q, Update):
        await update_or_q.message.reply_text(txt, reply_markup=InlineKeyboardMarkup(kb))
    else:
        q=update_or_q; await q.edit_message_text(txt, reply_markup=InlineKeyboardMarkup(kb))

async def menu_router(update:Update, context:ContextTypes.DEFAULT_TYPE):
    q=update.callback_query; await q.answer()
    _,_,action,pid = q.data.split(":"); pid=int(pid)
    if action=="addexp":
        context.user_data["exp"]={"pid":pid}; await q.edit_message_text("المسمى الوظيفي (Role):"); return EXP_ROLE
    if action=="addedu":
        context.user_data["edu"]={"pid":pid}; await q.edit_message_text("الدرجة العلمية:"); return EDU_DEGREE
    if action=="skills":
        context.user_data["skills_pid"]=pid; await q.edit_message_text("أرسل المهارات مفصولة بفواصل:"); return SKILLS_SET
    if action=="export":
        await q.edit_message_text("اختر طريقة التصدير / المعاينة:")
        return await show_export_menu(q, context, pid)

# --- Experience flow ---
async def exp_role(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["exp"]["role"]=update.message.text.strip()
    await update.message.reply_text("اسم الشركة:"); return EXP_COMPANY

async def exp_company(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["exp"]["company"]=update.message.text.strip()
    await update.message.reply_text("تاريخ البدء (مثال 01/2023):"); return EXP_START

async def exp_start(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["exp"]["start_date"]=update.message.text.strip()
    await update.message.reply_text("تاريخ الانتهاء (أو اكتب Present):"); return EXP_END

async def exp_end(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["exp"]["end_date"]=update.message.text.strip()
    await update.message.reply_text("أرسل نقاط الإنجاز (كل سطر نقطة، رسالة واحدة):"); return EXP_BULLETS

async def exp_bullets(update:Update, context:ContextTypes.DEFAULT_TYPE):
    lines=[l.strip("• ").strip() for l in update.message.text.splitlines() if l.strip()]
    e=context.user_data.get("exp",{})
    db.add_experience(e["pid"], e["company"], e["role"], e["start_date"], e["end_date"], lines)
    await update.message.reply_text("تمت إضافة الخبرة.")
    await show_menu(update, context, e["pid"]); return MENU

# --- Education flow ---
async def edu_degree(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["edu"]["degree"]=update.message.text.strip()
    await update.message.reply_text("التخصص:"); return EDU_MAJOR

async def edu_major(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["edu"]["major"]=update.message.text.strip()
    await update.message.reply_text("اسم الجامعة/المعهد:"); return EDU_SCHOOL

async def edu_school(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["edu"]["school"]=update.message.text.strip()
    await update.message.reply_text("سنة التخرج:"); return EDU_YEAR

async def edu_year(update:Update, context:ContextTypes.DEFAULT_TYPE):
    context.user_data["edu"]["year"]=update.message.text.strip()
    ed=context.user_data["edu"]
    db.add_education(ed["pid"], ed["degree"], ed["major"], ed["school"], ed["year"])
    await update.message.reply_text("تمت إضافة التعليم.")
    await show_menu(update, context, ed["pid"]); return MENU

# --- Skills ---
async def skills_set(update:Update, context:ContextTypes.DEFAULT_TYPE):
    pid=context.user_data.get("skills_pid")
    db.set_skills(pid, update.message.text.strip())
    await update.message.reply_text("تم تعيين المهارات.")
    await show_menu(update, context, pid); return MENU

# --- Export / Preview ---
async def show_export_menu(q, context:ContextTypes.DEFAULT_TYPE, pid:int):
    user_id=q.from_user.id
    buttons=[
        [InlineKeyboardButton("👀 معاينة (صورة)", callback_data=f"cv:export:preview:{pid}")],
        [InlineKeyboardButton("📄 تصدير DOCX", callback_data=f"cv:export:docx:{pid}")],
    ]
    if db.is_vip(user_id) or user_is_owner(q.from_user):
        buttons.append([InlineKeyboardButton("🧾 تصدير PDF (عالي الجودة)", callback_data=f"cv:export:pdf:{pid}")])
        buttons.append([InlineKeyboardButton("✉️ Cover Letter", callback_data=f"cv:export:cover:{pid}")])
    else:
        buttons.append([InlineKeyboardButton("⭐ ترقية إلى VIP", url=PAYLINK_UPGRADE_URL or "https://example.com")])
    await q.edit_message_reply_markup(reply_markup=InlineKeyboardMarkup(buttons))
    return CONFIRM_EXPORT

async def export_router(update:Update, context:ContextTypes.DEFAULT_TYPE):
    q=update.callback_query; await q.answer()
    _, _, kind, pid = q.data.split(":"); pid=int(pid)
    user_id=q.from_user.id

    if kind == "preview":
        await q.edit_message_text("جارٍ إنشاء معاينة…")
        try:
            html=render_html_for_profile(pid, db)
            try:
                png=await docraptor_convert(html, kind="png")  # قد يفشل على الخطة المجانية
                out=EXPORTS_DIR/f"preview_{pid}.png"; out.write_bytes(png)
                with open(out,"rb") as f:
                    await q.message.reply_photo(f, caption="هذه المعاينة. إذا مناسب اختر PDF أو DOCX.")
            except Exception as e:
                log.warning("PNG preview failed, falling back to PDF: %s", e)
                pdf=await docraptor_convert(html, kind="pdf")
                out=EXPORTS_DIR/f"preview_{pid}.pdf"; out.write_bytes(pdf)
                with open(out,"rb") as f:
                    await q.message.reply_document(InputFile(f, filename=out.name), caption="معاينة PDF")
        except Exception as e:
            await q.message.reply_text(f"تعذّرت المعاينة: {e}")
        await show_menu(q, context, pid); return MENU

    if kind=="docx":
        is_owner=user_is_owner(q.from_user)
        if not (db.is_vip(user_id) or is_owner) and (not db.free_once_available(user_id)):
            await q.edit_message_text("استخدمت محاولتك المجانية الوحيدة. رجاءً قم بالترقية إلى VIP.")
            return ConversationHandler.END
        await q.edit_message_text("جارٍ إنشاء DOCX…")
        path=render_docx_for_profile(pid, db)
        if not (db.is_vip(user_id) or is_owner): db.mark_free_once_used(user_id)
        with open(path,"rb") as f:
            await q.message.reply_document(InputFile(f, filename=path.name), caption="تم إنشاء السيرة ✨")
        await show_menu(q, context, pid); return MENU

    if kind=="pdf":
        if not (db.is_vip(user_id) or user_is_owner(q.from_user)):
            await q.edit_message_text("ميزة PDF لعملاء VIP فقط.")
            return ConversationHandler.END
        await q.edit_message_text("جارٍ إنشاء PDF…")
        try:
            html=render_html_for_profile(pid, db)
            pdf=await docraptor_convert(html, kind="pdf")
            out=EXPORTS_DIR/f"cv_{pid}.pdf"; out.write_bytes(pdf)
            with open(out,"rb") as f:
                await q.message.reply_document(InputFile(f, filename=out.name), caption="PDF جاهز ✅")
        except Exception as e:
            await q.message.reply_text(f"فشل توليد PDF: {e}\nسأرسل DOCX بدلًا منه.")
            path=render_docx_for_profile(pid, db)
            with open(path,"rb") as f:
                await q.message.reply_document(InputFile(f, filename=path.name))
        await show_menu(q, context, pid); return MENU

    if kind=="cover":
        if not (db.is_vip(user_id) or user_is_owner(q.from_user)):
            await q.edit_message_text("Cover Letter لعملاء VIP فقط."); return ConversationHandler.END
        profile, exps, edus, skills = db.fetch_full_profile(pid)
        lang=profile.get("lang","ar")
        body = (
            f"السادة المحترمون،\n\n"
            f"أتقدم لوظيفة {profile.get('title','')} ولدي خبرات ذات صلة.\n"
            f"أرفقت سيرتي الذاتية وأتطلع لفرصة مقابلة.\n\n"
            f"تحياتي،\n{profile.get('full_name','')}\n{profile.get('phone','')} • {profile.get('email','')}"
        ) if lang=="ar" else (
            f"Dear Hiring Team,\n\nI am applying for the {profile.get('title','')} role. "
            f"Please find my resume attached. I would welcome the opportunity to discuss my fit.\n\n"
            f"Kind regards,\n{profile.get('full_name','')}\n{profile.get('phone','')} • {profile.get('email','')}"
        )
        fn=EXPORTS_DIR/f"cover_{pid}.txt"; fn.write_text(body, encoding="utf-8")
        with open(fn,"rb") as f:
            await q.message.reply_document(InputFile(f, filename=fn.name), caption="Cover Letter")
        await show_menu(q, context, pid); return MENU

# -------------- Mini HTTP server (/health) --------------
async def create_app_and_site(app_tg: Application):
    async def root(request): return web.Response(text="OK")
    async def health(request):
        return web.json_response({"ok": True, "service": "cvbot", "time": datetime.utcnow().isoformat()})
    app=web.Application()
    # مهم: سجّل GET فقط (HEAD يتولد تلقائيًا) — لا تسجّل web.head(..) حتى لا يظهر خطأ "method HEAD is already registered"
    app.add_routes([web.get("/",root), web.get("/health",health)])
    runner=web.AppRunner(app); await runner.setup()
    site=web.TCPSite(runner, host="0.0.0.0", port=PORT); await site.start()
    log.info("aiohttp listening on :%s", PORT)

# -------------- Main --------------
async def _post_init(app:Application):
    await set_my_commands(app)
    await create_app_and_site(app)

def main():
    token=os.getenv("BOT_TOKEN","")
    if not token: raise RuntimeError("BOT_TOKEN is missing")
    application=Application.builder().token(token).post_init(_post_init).build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_cmd))
    application.add_handler(CommandHandler("upgrade", upgrade_cmd))

    cv_conv=ConversationHandler(
        entry_points=[CommandHandler("cv", cv_entry)],
        states={
            ASK_LANG:[CallbackQueryHandler(cv_set_lang, pattern=r"^cv:lang:")],
            ASK_TPL:[CallbackQueryHandler(cv_set_tpl, pattern=r"^cv:tpl:")],
            ASK_NAME:[MessageHandler(filters.TEXT & ~filters.COMMAND, cv_name)],
            ASK_TITLE:[MessageHandler(filters.TEXT & ~filters.COMMAND, cv_title)],
            ASK_PHONE:[MessageHandler(filters.TEXT & ~filters.COMMAND, cv_phone)],
            ASK_EMAIL:[MessageHandler(filters.TEXT & ~filters.COMMAND, cv_email)],
            ASK_CITY:[MessageHandler(filters.TEXT & ~filters.COMMAND, cv_city)],
            ASK_LINKS:[MessageHandler(filters.TEXT & ~filters.COMMAND, cv_links)],
            ASK_SUMMARY:[MessageHandler(filters.TEXT & ~filters.COMMAND, cv_summary)],
            MENU:[CallbackQueryHandler(menu_router, pattern=r"^cv:menu:")],
            EXP_ROLE:[MessageHandler(filters.TEXT & ~filters.COMMAND, exp_role)],
            EXP_COMPANY:[MessageHandler(filters.TEXT & ~filters.COMMAND, exp_company)],
            EXP_START:[MessageHandler(filters.TEXT & ~filters.COMMAND, exp_start)],
            EXP_END:[MessageHandler(filters.TEXT & ~filters.COMMAND, exp_end)],
            EXP_BULLETS:[MessageHandler(filters.TEXT & ~filters.COMMAND, exp_bullets)],
            EDU_DEGREE:[MessageHandler(filters.TEXT & ~filters.COMMAND, edu_degree)],
            EDU_MAJOR:[MessageHandler(filters.TEXT & ~filters.COMMAND, edu_major)],
            EDU_SCHOOL:[MessageHandler(filters.TEXT & ~filters.COMMAND, edu_school)],
            EDU_YEAR:[MessageHandler(filters.TEXT & ~filters.COMMAND, edu_year)],
            SKILLS_SET:[MessageHandler(filters.TEXT & ~filters.COMMAND, skills_set)],
            CONFIRM_EXPORT:[CallbackQueryHandler(export_router, pattern=r"^cv:export:")],
        },
        fallbacks=[],
        name="cv_conv",
        persistent=False,
    )
    application.add_handler(cv_conv)

    log.info("Bot starting…")
    application.run_polling()

if __name__ == "__main__":
    main()

